import pandas as pd
import re
import os
import sys
from difflib import get_close_matches

import pdfplumber
from docx import Document

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    full_text = ""

    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                full_text += page.extract_text() + "\n"
                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row and len(row) >= 2:
                            key, val = row[0], row[1]
                            full_text += f"{key.strip()} : {val.strip()}\n"

    elif ext == ".docx":
        from docx import Document
        doc = Document(file_path)

        # Paragraphs
        full_text += "\n".join([para.text for para in doc.paragraphs if para.text.strip()]) + "\n"

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    key = cells[0].text.strip()
                    val = cells[1].text.strip()
                    full_text += f"{key} : {val}\n"

    else:
        raise ValueError("Unsupported file format: " + ext)

    return full_text


def main(doc_path, excel_path="KPITemplate.xlsx"):
    print("[+] Loading KPI Excel...")
    df = pd.read_excel(excel_path, sheet_name="Sheet1", skiprows=1)

    print(f"[+] Extracting text from: {doc_path}")
    text = extract_text_from_file(doc_path)

    print("[+] Processing and matching metrics...")

    kpi_rules = {
        "application owner": lambda txt: re.search(r'Point of contact\s*\(App Team\)\s*[:\-]?\s*(.+)', txt, re.IGNORECASE),
        "user roles": lambda txt: re.search(r'user roles\s+(\d+)', txt, re.IGNORECASE),
        "testing environment": lambda txt: re.search(r'testing environment\s+(\w+)', txt, re.IGNORECASE),
        "tools used": lambda txt: re.findall(r'(AppScan|Burp Suite)', txt, re.IGNORECASE),
        "vulnerabilities": lambda txt: re.search(r'identified.*?(\d+)\s+vulnerabilities.*?(\d+)\s+false', txt, re.IGNORECASE),
    }

    for idx, row in df.iterrows():
        metric_text = str(row.iloc[2])
        if pd.isna(metric_text):
            continue

        best_match = get_close_matches(metric_text.lower(), kpi_rules.keys(), n=1, cutoff=0.4)
        if not best_match:
            continue

        key = best_match[0]
        result = kpi_rules[key](text)

        if result:
            if isinstance(result, list):
                comment = ", ".join(set(result))
            elif isinstance(result, re.Match):
                comment = " | ".join(g for g in result.groups() if g)
            else:
                comment = str(result)

            df.at[idx, "Responses"] = "Completed"
            df.at[idx, "Additional Comments (If any)"] = comment

    out_file = "KPITemplate_filled.xlsx"
    df.to_excel(out_file, index=False)
    print(f"[✓] KPI checklist updated and saved to: {out_file}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python fill_kpi.py <document.pdf/docx>")
    else:
        main(sys.argv[1])

